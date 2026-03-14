#!/usr/bin/env python3
"""
PDF to Word Converter - Converts PDFs to formatted Word documents

Usage:
    python tools/pdf_to_word.py --input input.pdf --output output.docx [--ocr]

Example:
    python tools/pdf_to_word.py --input document.pdf --output formatted.docx --ocr
"""

import os
import sys
import argparse
import re
import base64
import io
from pathlib import Path
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Import PDF processing libraries
try:
    import pymupdf4llm
    import fitz  # PyMuPDF
except ImportError:
    print("Error: PyMuPDF libraries not installed. Run: pip install pymupdf4llm pymupdf", file=sys.stderr)
    sys.exit(1)

# Import Word generation library
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)

# Import image processing
try:
    from PIL import Image
except ImportError:
    print("Error: Pillow not installed. Run: pip install pillow", file=sys.stderr)
    sys.exit(1)

# Import Anthropic for image description
try:
    import anthropic
except ImportError:
    print("Error: anthropic not installed. Run: pip install anthropic", file=sys.stderr)
    sys.exit(1)


def clean_filename(filename):
    """
    Clean PDF filename according to specifications:
    - Remove leading numbers (e.g., "058-" or "23_")
    - Trim to 50 characters maximum
    - Remove cut-off words at the end
    - Keep hyphens and proper formatting

    Example: "058-texas-adapted-genetic-strategies-for-beef-cattle-v.pdf"
          -> "texas-adapted-genetic-strategies-for-beef-cattle.pdf"
    """
    # Remove file extension
    name_without_ext = Path(filename).stem

    # Remove leading numbers and separators
    cleaned = re.sub(r'^\d+[-_]', '', name_without_ext)

    # If longer than 50 chars, trim at word boundary
    if len(cleaned) > 50:
        # Trim to 50 chars
        cleaned = cleaned[:50]
        # Find last hyphen or underscore to avoid cutting words
        last_sep = max(cleaned.rfind('-'), cleaned.rfind('_'), cleaned.rfind(' '))
        if last_sep > 30:  # Only trim if we're not losing too much
            cleaned = cleaned[:last_sep]

    # Remove any trailing incomplete words (single letters at end after separator)
    cleaned = re.sub(r'[-_]\w$', '', cleaned)

    # Add .pdf extension back
    return cleaned + '.pdf'


def extract_authors(text_lines):
    """
    Extract and clean author names from the beginning of the document.
    - Typically appears under the title
    - Remove position titles
    - Remove superscripts and special characters
    - Join with commas
    """
    authors = []

    # Look for author patterns in first 10 lines (after title)
    for line in text_lines[1:10]:  # Skip first line (likely title)
        line = line.strip()

        # Skip empty lines
        if not line:
            continue

        # Check if line looks like authors (contains names, not all caps unless abbreviated)
        # Authors typically have capitals, commas, and possibly "and"
        if re.search(r'[A-Z][a-z]+', line) and not line.isupper():
            # Remove superscript numbers and symbols
            clean_line = re.sub(r'[⁰¹²³⁴⁵⁶⁷⁸⁹\*†‡§¶]', '', line)

            # Remove common title words
            title_words = ['Ph\.?D\.?', 'M\.?D\.?', 'Professor', 'Dr\.', 'Researcher',
                          'Specialist', 'Extension', 'Assistant', 'Associate', 'Department']
            for title in title_words:
                clean_line = re.sub(title, '', clean_line, flags=re.IGNORECASE)

            # Clean up extra spaces
            clean_line = re.sub(r'\s+', ' ', clean_line).strip()

            if clean_line:
                authors.append(clean_line)

        # Stop after finding authors (before main content starts)
        if len(authors) > 0 and (line.isupper() or len(line) > 100):
            break

    return ', '.join(authors) if authors else None


def fix_heading_case(text):
    """
    Fix all-caps headings to title case.
    Only affects lines that are entirely uppercase and reasonably short (headings).
    """
    lines = text.split('\n')
    fixed_lines = []

    for line in lines:
        stripped = line.strip()
        # If line is all caps and looks like a heading (short, no punctuation at end)
        if stripped.isupper() and len(stripped) < 100 and not stripped.endswith('.'):
            # Convert to title case
            fixed_lines.append(stripped.title())
        else:
            fixed_lines.append(line)

    return '\n'.join(fixed_lines)


def remove_special_bullets(text):
    """
    Remove non-circular/non-disc bullets.
    Common problematic bullets: arrows, squares, custom symbols.
    """
    # Remove various bullet symbols, keeping standard bullets (•, ○, -)
    # Replace with standard bullet point
    text = re.sub(r'[▪▫■□●○◆◇♦►▻→➔➢]', '•', text)

    # Remove bullets with dashes attached to words
    text = re.sub(r'—(\w)', r' \1', text)  # em-dash before word
    text = re.sub(r'–(\w)', r' \1', text)  # en-dash before word

    return text


def fix_fractions(text):
    """
    Fix fractional number formatting.
    Correct format: "1 1/8 inches" or "1/8 inches"
    """
    # Fix fractions without spaces: "1-1/8" -> "1 1/8"
    text = re.sub(r'(\d)-(\d)/(\d)', r'\1 \2/\3', text)

    # Ensure space before "inches" or other units
    text = re.sub(r'(\d/\d)(\w)', r'\1 \2', text)

    return text


def cleanup_text(text):
    """
    Apply all text cleanup rules to extracted text.
    """
    # Fix all-caps headings
    text = fix_heading_case(text)

    # Remove extra line breaks (3+ newlines -> 2)
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Remove double spacing
    text = re.sub(r'  +', ' ', text)

    # Remove non-circular bullets
    text = remove_special_bullets(text)

    # Fix fractions
    text = fix_fractions(text)

    return text


def describe_image_with_claude(image_data, figure_number):
    """
    Use Claude's vision API to describe an image and determine if it's a figure or table.
    Returns dict with 'type' (Figure/Table), 'description', and 'suggested_name'.
    """
    api_key = os.getenv('ANTHROPIC_API_KEY')
    if not api_key or api_key == 'your_anthropic_key_here':
        print("Warning: ANTHROPIC_API_KEY not configured. Skipping image description.", file=sys.stderr)
        return {
            'type': 'Figure',
            'description': 'Image',
            'suggested_name': f'{figure_number:03d} Figure {figure_number} Image'
        }

    try:
        client = anthropic.Anthropic(api_key=api_key)

        # Encode image to base64
        buffered = io.BytesIO()
        image_data.save(buffered, format="PNG")
        img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')

        # Ask Claude to describe the image
        message = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=300,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": "image/png",
                            "data": img_base64
                        }
                    },
                    {
                        "type": "text",
                        "text": """Analyze this image from a PDF document and provide:

1. TYPE: Determine if this is a "Figure" or "Table"
   - Figure: photos, illustrations, charts, graphs, diagrams, drawings
   - Table: structured data in rows and columns with headers

2. DESCRIPTION: Provide a SPECIFIC, DESCRIPTIVE name (3-6 words) that captures the key content
   - For Tables: describe what data is shown (e.g., "Production Data by Region", "Cost Analysis Summary")
   - For Figures: describe what is depicted (e.g., "Cattle Grazing in Pasture", "Farm Equipment Diagram")
   - Be SPECIFIC and DESCRIPTIVE, not generic

Respond in this exact format:
TYPE: [Figure or Table]
DESCRIPTION: [specific descriptive name]

Examples of GOOD descriptions:
TYPE: Table
DESCRIPTION: Monthly Revenue Breakdown

TYPE: Figure
DESCRIPTION: Holstein Cows in Feedlot

TYPE: Figure
DESCRIPTION: Crop Rotation Schedule Chart

Examples of BAD descriptions (too generic):
- "Image"
- "Data"
- "Chart"
- "Photo"

Be specific and descriptive!"""
                    }
                ]
            }]
        )

        # Parse Claude's response
        response_text = message.content[0].text.strip()

        # Debug: Print Claude's response
        print(f"    Claude response: {response_text[:100]}...")

        # Extract type and description
        type_match = re.search(r'TYPE:\s*(Figure|Table)', response_text, re.IGNORECASE)
        desc_match = re.search(r'DESCRIPTION:\s*(.+)', response_text, re.IGNORECASE)

        img_type = type_match.group(1).title() if type_match else 'Figure'
        description = desc_match.group(1).strip() if desc_match else 'Image'

        # Debug: Print extracted values
        print(f"    Extracted - Type: {img_type}, Description: {description}")

        # Generate suggested filename
        suggested_name = f'{figure_number:03d} {img_type} {figure_number} {description}'

        return {
            'type': img_type,
            'description': description,
            'suggested_name': suggested_name
        }

    except Exception as e:
        print(f"Warning: Could not describe image with Claude: {e}", file=sys.stderr)
        return {
            'type': 'Figure',
            'description': 'Image',
            'suggested_name': f'{figure_number:03d} Figure {figure_number} Image'
        }


def extract_and_save_images(pdf_path, output_dir):
    """
    Extract all images from PDF, describe them with Claude, and save with descriptive names.
    Naming format: NNN Type N Description
    - NNN = overall order in PDF (001, 002, 003...)
    - Type = Figure or Table
    - N = type-specific counter (Figure 1, Figure 2... / Table 1, Table 2...)
    Returns list of dicts with image info: {path, page, type, description, name}
    """
    print(f"Extracting images from PDF...")

    # Create output directory if it doesn't exist
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    extracted_images = []

    try:
        doc = fitz.open(pdf_path)

        # Two separate counters
        overall_number = 1  # Overall sequence in PDF (001, 002, 003...)
        figure_count = 0    # Count of figures only
        table_count = 0     # Count of tables only

        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images(full=True)

            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]

                    # Convert to PIL Image
                    pil_image = Image.open(io.BytesIO(image_bytes))

                    # Skip small images (likely icons, logos, or decorations)
                    # Use a higher threshold to avoid capturing branding elements
                    # Also skip very wide/short images (likely headers/footers)
                    aspect_ratio = pil_image.width / pil_image.height if pil_image.height > 0 else 0

                    if pil_image.width < 300 or pil_image.height < 150:
                        print(f"  Skipping small image on page {page_num + 1} ({pil_image.width}x{pil_image.height})")
                        continue

                    if aspect_ratio > 4:  # Wide images (>4:1) are likely headers/banners
                        print(f"  Skipping header/banner image on page {page_num + 1} ({pil_image.width}x{pil_image.height}, ratio: {aspect_ratio:.1f})")
                        continue

                    # Describe image with Claude (temporarily pass overall_number for display)
                    print(f"  Analyzing image {overall_number} on page {page_num + 1}...")
                    description_info = describe_image_with_claude(pil_image, overall_number)

                    # Determine type and increment appropriate counter
                    img_type = description_info['type']
                    if img_type == 'Figure':
                        figure_count += 1
                        type_number = figure_count
                    else:  # Table
                        table_count += 1
                        type_number = table_count

                    # Generate filename: NNN Type N Description
                    description = description_info['description']
                    safe_desc = re.sub(r'[^\w\s-]', '', description)
                    safe_desc = re.sub(r'[-\s]+', ' ', safe_desc).strip()

                    filename = f"{overall_number:03d} {img_type} {type_number} {safe_desc}.{image_ext}"

                    # Save image
                    image_path = output_dir / filename
                    pil_image.save(image_path)

                    print(f"  ✓ Saved: {filename}")

                    extracted_images.append({
                        'path': str(image_path),
                        'filename': filename,
                        'page': page_num + 1,
                        'type': img_type,
                        'type_number': type_number,
                        'overall_number': overall_number,
                        'description': description
                    })

                    overall_number += 1

                except Exception as e:
                    print(f"  Error extracting image on page {page_num + 1}: {e}", file=sys.stderr)
                    continue

        doc.close()
        print(f"✓ Extracted {len(extracted_images)} images ({figure_count} figures, {table_count} tables)")
        return extracted_images

    except Exception as e:
        print(f"Error extracting images: {e}", file=sys.stderr)
        return []


def extract_pdf_content(pdf_path, use_ocr=False, extract_images_flag=True):
    """
    Extract text, structure, and images from PDF using PyMuPDF4LLM.
    Returns dict with extracted content and metadata.
    """
    print(f"Extracting content from: {pdf_path}")

    try:
        # Use PyMuPDF4LLM for structure-aware extraction
        md_text = pymupdf4llm.to_markdown(pdf_path)

        # Also open with PyMuPDF for additional metadata and images
        doc = fitz.open(pdf_path)

        # Extract text blocks with position info
        all_blocks = []
        for page_num, page in enumerate(doc):
            blocks = page.get_text("dict")["blocks"]
            all_blocks.extend([(page_num, block) for block in blocks])

        # Count images for reference
        image_count = 0
        for page_num, page in enumerate(doc):
            image_list = page.get_images()
            image_count += len(image_list)

        doc.close()

        # Extract and save images with descriptions
        extracted_images = []
        if extract_images_flag and image_count > 0:
            # Create temp directory for images
            temp_image_dir = Path('.tmp/images')
            extracted_images = extract_and_save_images(pdf_path, temp_image_dir)

        return {
            'markdown_text': md_text,
            'blocks': all_blocks,
            'extracted_images': extracted_images,
            'success': True
        }

    except Exception as e:
        print(f"Error extracting PDF content: {e}", file=sys.stderr)
        return {'success': False, 'error': str(e)}


def add_comment_to_paragraph(paragraph, comment_text):
    """
    Add a comment to a paragraph in the Word document.
    This is a workaround since python-docx doesn't directly support comments.
    We'll add it as a note in brackets for now.
    """
    # For simplicity, we'll add comments as text in brackets
    # A full comment implementation would require more complex XML manipulation
    paragraph.add_run(f" [{comment_text}]").font.italic = True


def detect_heading_level(line, prev_line=''):
    """
    Detect if a line is a heading and what level (H2, H3, H4).
    PyMuPDF4LLM marks headings with # symbols in markdown.
    """
    line = line.strip()

    # Check for markdown-style headings
    if line.startswith('###'):
        return 4, line.replace('###', '').strip()
    elif line.startswith('##'):
        return 3, line.replace('##', '').strip()
    elif line.startswith('#'):
        return 2, line.replace('#', '').strip()

    # Check for all-caps lines (likely headings)
    if line.isupper() and 5 < len(line) < 100 and not line.endswith('.'):
        return 3, line  # Default to H3 for all-caps

    return None, line


def generate_word_doc(content, output_path, original_filename):
    """
    Generate a Word document from extracted PDF content.
    Applies all formatting rules and structure.
    All body text will have consistent font, size, and color.
    All headings will have comments indicating their level.
    """
    print("Generating Word document...")

    doc = Document()

    # Get the markdown text
    md_text = content['markdown_text']
    lines = md_text.split('\n')

    # Clean filename for title
    cleaned_filename = clean_filename(original_filename)
    print(f"Cleaned filename: {cleaned_filename}")

    # Add title (first non-empty line or use filename)
    title_text = next((line.strip() for line in lines if line.strip()), cleaned_filename)
    if title_text.startswith('#'):
        title_text = title_text.lstrip('#').strip()

    title = doc.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Add H1 comment to title
    title.add_run(" [H1]").font.italic = True
    print(f"Added title: '{title_text}' - H1")

    # Extract and add authors
    authors = extract_authors(lines)
    if authors:
        author_para = doc.add_paragraph(authors)
        # Ensure consistent formatting for author text
        for run in author_para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.italic = True
        print(f"Added authors: {authors}")

    # Process remaining content
    in_list = False
    first_paragraph_done = False
    first_figure_placed = False

    for i, line in enumerate(lines[1:], 1):  # Skip first line (title)
        line = line.strip()

        if not line:
            continue

        # Skip author lines (already processed)
        if authors and line in authors:
            continue

        # Detect heading level
        heading_level, heading_text = detect_heading_level(line, lines[i-1] if i > 0 else '')

        if heading_level:
            # Add heading
            heading = doc.add_heading(heading_text, level=heading_level)

            # Add comment noting heading level
            comment_text = f"H{heading_level}"
            heading.add_run(f" [{comment_text}]").font.italic = True
            print(f"  Heading detected: '{heading_text}' - {comment_text}")
            in_list = False

        elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
            # Bullet point
            bullet_text = line.lstrip('•-* ').strip()
            bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
            # Ensure consistent formatting for bullet text
            for run in bullet_para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
            in_list = True

        elif line.startswith('[') and line.endswith(']'):
            # Likely an image reference from markdown - add placeholder
            para = doc.add_paragraph()
            run = para.add_run(f"[Figure placeholder: {line}]")
            run.font.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

            # Ensure first figure is after first paragraph
            if not first_figure_placed and not first_paragraph_done:
                print(f"  Warning: First figure should appear after first paragraph")
            first_figure_placed = True

        else:
            # Regular paragraph
            # Clean the text
            cleaned_line = cleanup_text(line)
            para = doc.add_paragraph(cleaned_line)
            # Ensure consistent formatting: same font, size, and color for all body text
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                # Color is already black by default
            in_list = False

            if not first_paragraph_done:
                first_paragraph_done = True

    # Note about extracted images
    if content.get('extracted_images'):
        print(f"\nExtracted images summary:")
        for img in content['extracted_images']:
            print(f"  - {img['filename']} (Page {img['page']})")

        # Add note about where images were saved
        note_para = doc.add_paragraph("\n[NOTE: Images have been extracted and saved to the .tmp/images folder. They will be uploaded to Google Drive Processed folder.]")
        for run in note_para.runs:
            run.font.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

        # List all extracted images with their descriptions
        doc.add_paragraph("\nExtracted Images:").runs[0].bold = True
        for img in content['extracted_images']:
            img_para = doc.add_paragraph(f"• {img['filename']} (Page {img['page']})")
            for run in img_para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)

    # Save the document
    doc.save(output_path)
    print(f"✓ Word document saved to: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='Convert PDF to formatted Word document')
    parser.add_argument('--input', required=True, help='Input PDF file path')
    parser.add_argument('--output', required=True, help='Output Word document path')
    parser.add_argument('--ocr', action='store_true', help='Use OCR for scanned documents')
    parser.add_argument('--extract-images', action='store_true', default=True,
                       help='Extract and describe images (default: True)')
    parser.add_argument('--no-extract-images', action='store_false', dest='extract_images',
                       help='Skip image extraction')

    args = parser.parse_args()

    # Validate input file exists
    if not Path(args.input).exists():
        print(f"Error: Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    # Extract content from PDF
    content = extract_pdf_content(args.input, use_ocr=args.ocr,
                                  extract_images_flag=args.extract_images)

    if not content['success']:
        print(f"Failed to extract PDF content: {content.get('error', 'Unknown error')}", file=sys.stderr)
        sys.exit(1)

    # Generate Word document
    original_filename = Path(args.input).name
    generate_word_doc(content, args.output, original_filename)

    print("\n✓ Conversion complete!")
    print(f"  Input:  {args.input}")
    print(f"  Output: {args.output}")

    # Report extracted images
    if content.get('extracted_images'):
        print(f"  Images: {len(content['extracted_images'])} extracted to .tmp/images/")
        return content['extracted_images']

    return []


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\nInterrupted by user", file=sys.stderr)
        sys.exit(130)
    except Exception as e:
        print(f"Unexpected error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        sys.exit(1)
